import tkinter
import tkinter.messagebox
import customtkinter
import random
import time
import datetime
from docx import Document
from docx.shared import Pt, Cm
from babel.dates import format_date, format_datetime, format_time
from tkcalendar import DateEntry
import numpy as np

from docx2pdf import convert

customtkinter.set_appearance_mode("Dark")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("green")  # Themes: "blue" (standard), "green", "dark-blue"

resolution = ['Rooms', 'Kitchen & Cabinets', 'Stairs', 'Ceiling, Windows & Doors']
hardnessArray = [4, 2, 1, 2]
hardness = []

d = {}
o = []
h = 0
g = 0

weekends = 0
# weekendsProduction = 0

schedule = 1

total = []
types = []


class App(customtkinter.CTk):
    WIDTH = 800
    HEIGHT = 800

    def __init__(self):
        super().__init__()

        self.title("BM Modul Sekvens Selektion")
        self.geometry(f"{App.WIDTH}x{App.HEIGHT}")
        self.protocol("WM_DELETE_WINDOW", self.on_closing)  # call .on_closing() when app gets closed

        self.label_9 = customtkinter.CTkLabel(master=self, text="MODULES",
                                              text_font=("Samsung Sharp Sans", -20))  # font name and size in px
        self.label_9.place(anchor='c', x=245, y=390, width=100)

        self.label_9 = customtkinter.CTkLabel(master=self, text="PRODUCTION SEQUENCE",
                                              text_font=("Samsung Sharp Sans", -20))  # font name and size in px
        self.label_9.place(anchor='c', x=555, y=390, width=250)

        self.orderField = customtkinter.CTkTextbox(master=self, state='disabled')
        self.orderField.place(anchor='nw', x=100, y=410, width=290, height=350)

        self.sequenceField = customtkinter.CTkTextbox(master=self, state='disabled')
        self.sequenceField.place(anchor='nw', x=410, y=410, width=290, height=350)

        self.label_1 = customtkinter.CTkLabel(master=self, text="BM Byggeindustri A/S",
                                              text_font=("Samsung Sharp Sans", -30))  # font name and size in px
        self.label_1.place(anchor='c', x=400, y=100)

        self.label_2 = customtkinter.CTkLabel(master=self, text="Production Planning Software v1.3",
                                              text_font=("Samsung Sharp Sans", -15))  # font name and size in px
        self.label_2.place(anchor='c', x=400, y=140)

        self.entryVar = customtkinter.StringVar()
        self.entryVar.set(value=1)
        self.entryVar.trace_add("write", self.entryVarLockout)

        self.blockVar = customtkinter.StringVar()
        self.blockVar.set(value='01')

        self.floorVar = customtkinter.StringVar()
        self.floorVar.set(value='00')

        self.sequenceVar = customtkinter.StringVar()
        self.sequenceVar.set(value='1')

        self.amountVar = customtkinter.StringVar()
        self.amountVar.set(value=0)

        self.label_8 = customtkinter.CTkLabel(master=self, textvariable=self.amountVar,
                                              text_font=("Samsung Sharp Sans", -10))  # font name and size in px
        self.label_8.place(anchor='c', x=500, y=275, width=50)

        self.label_3 = customtkinter.CTkLabel(master=self, text="AMOUNT",
                                              text_font=("Samsung Sharp Sans", -10))  # font name and size in px
        self.label_3.place(anchor='c', x=270, y=275, width=50)

        self.label_3 = customtkinter.CTkLabel(master=self, text="BLOCK",
                                              text_font=("Samsung Sharp Sans", -10))  # font name and size in px
        self.label_3.place(anchor='c', x=505, y=275, width=50)

        self.label_4 = customtkinter.CTkLabel(master=self, text="FLOOR",
                                              text_font=("Samsung Sharp Sans", -10))  # font name and size in px
        self.label_4.place(anchor='c', x=560, y=275, width=50)

        self.label_5 = customtkinter.CTkLabel(master=self, text="NUMBER",
                                              text_font=("Samsung Sharp Sans", -10))  # font name and size in px
        self.label_5.place(anchor='c', x=615, y=275, width=50)

        self.entry = customtkinter.CTkEntry(master=self, placeholder_text="Amount", textvariable=self.entryVar)
        self.entry.place(anchor='c', x=270, y=300, width=100)

        self.blockEntry = customtkinter.CTkEntry(master=self, placeholder_text="Amount", textvariable=self.blockVar)
        self.blockEntry.place(anchor='c', x=505, y=300, width=50)

        self.floorEntry = customtkinter.CTkEntry(master=self, placeholder_text="Amount", textvariable=self.floorVar)
        self.floorEntry.place(anchor='c', x=560, y=300, width=50)

        self.sequenceEntry = customtkinter.CTkEntry(master=self, placeholder_text="Amount",
                                                    textvariable=self.sequenceVar)
        self.sequenceEntry.place(anchor='c', x=615, y=300, width=50)

        self.selector_1 = customtkinter.StringVar()
        self.selector_1.trace_add('write', self.selectorLockout1)

        self.selector_2 = customtkinter.StringVar()
        self.selector_2.trace_add('write', self.selectorLockout2)

        self.switch_1 = customtkinter.CTkSwitch(master=self, text="Contains First Module", onvalue=True, offvalue=False,
                                                variable=self.selector_1)
        self.switch_1.place(anchor='w', x=40, y=300)

        self.switch_2 = customtkinter.CTkSwitch(master=self, text="Contains Final Module", onvalue=1, offvalue=0,
                                                variable=self.selector_2)
        self.switch_2.place(anchor='w', x=40, y=340)

        self.button_1 = customtkinter.CTkButton(master=self, text="Add Module", command=self.button_event)
        self.button_1.place(anchor='c', x=400, y=300)

        self.button_2 = customtkinter.CTkButton(master=self, text="Calculate Sequence", command=self.button_event2,
                                                fg_color=("red"), hover_color=("darkred"))
        self.button_2.place(anchor='c', x=400, y=340)

        self.optionlist = customtkinter.StringVar()
        self.optionlist.trace_add("write", self.optionRetreive)
        self.options = []

        self.selector_3 = customtkinter.CTkComboBox(master=self, variable=self.optionlist, values=self.options)
        self.selector_3.place(anchor='c', x=270, y=340, width=100)

        ## Production date selector
        now = datetime.datetime.now()
        year = datetime.datetime.strftime(now, '%Y')
        month = datetime.datetime.strftime(now, '%m')
        day = datetime.datetime.strftime(now, '%d')

        self.cal = DateEntry(self, width=12, year=int(year), month=int(month), day=int(day),
                             background='green', foreground='white', borderwidth=2)
        self.cal.place(anchor='c', x=560, y=340, width=160, height=30)

        for k in range(len(hardnessArray)):
            hardness = []
            for j in range(hardnessArray[k]):
                if hardnessArray[k] == 1:
                    hardness.append(str(0))
                hardness.append(str(j + 1))
            o.append(hardness)

        print(o[1])

        for i in range(len(resolution)):
            var = "hardness" + str(i)
            d["stringVar{0}".format(i)] = customtkinter.StringVar(self, 1, var)

        locals().update(d)
        for i in range(len(resolution)):
            values = o[i]
            print(values)
            self.selector = customtkinter.CTkComboBox(master=self, variable=d.get("stringVar" + str(i)), values=values)

            # Calculate the optimal position for the dropdown...
            center = App.WIDTH / 2
            width = 50
            iPos = (center / 2) + (width / 2) + ((575 - 225) / (len(resolution) - 1)) * i

            self.selector.place(anchor='c', x=iPos, y=230, width=50)

            self.label = customtkinter.CTkLabel(master=self, text=resolution[i], text_font=("Samsung Sharp Sans", -14))
            self.label.place(anchor='c', x=iPos, y=200, width=200)

    def button_event(self):
        global h

        try:
            integer_result = int(self.entryVar.get())
        except ValueError:
            None
        else:
            for g in range(integer_result):
                moduleCode = []
                moduleCode.append(self.blockVar.get())
                moduleCode.append(self.floorVar.get())
                moduleCode.append(self.sequenceVar.get())

                choices = ()
                for j in range(len(resolution)):
                    choices = choices + (int((d.get("stringVar"+str(j)).get())),)
                    #choices = choices + (random.randint(1, 5),)
                moduleTypes = choices

                if self.selector_1.get() == '1':
                    choices = choices + ('End',)
                elif self.selector_2.get() == '1':
                    choices = choices + ('First',)
                else:
                    choices = choices + ('',)

                self.orderField.configure(state='normal')
                self.orderField.insert(tkinter.END,
                                       str(choices[:len(resolution)]) + '  ->  ' + str(moduleCode)+'\n')
                self.orderField.configure(state='disabled')

                choices = choices + ('m' + str(h),)
                choices = choices + (moduleCode,)
                print(choices)
                total.append(choices)
                self.sequenceVar.set(int(self.sequenceVar.get()) + 1)

        if moduleTypes not in types:
            self.options.append("Type " + str(h))
            self.selector_3.configure(values=self.options)
            types.append(moduleTypes)
            h = h + 1

    def selectorLockout1(self, var, index, mode):
        if self.selector_2.get() == '1':
            self.switch_2.deselect()

    def selectorLockout2(self, var, index, mode):
        if self.selector_1.get() == '1':
            self.switch_1.deselect()

    def optionRetreive(self, var, index, mode):
        try:
            integer_result = int(self.optionlist.get()[5:])
        except ValueError:
            None
        else:
            options = types[integer_result]
            for j in range(len(resolution)):
                d.get("stringVar" + str(j)).set(options[j])

    def entryVarLockout(self, var, index, mode):
        try:
            integer_result = int(self.entryVar.get())
        except ValueError:
            # print("not a valid integer")
            self.blockEntry.configure(state=tkinter.DISABLED)
            self.floorEntry.configure(state=tkinter.DISABLED)
            self.sequenceEntry.configure(state=tkinter.DISABLED)
            # self.blockVar.set('')
            # self.floorVar.set('')
            # self.sequenceVar.set('')
        else:
            # print("unlocked")
            self.blockEntry.configure(state=tkinter.NORMAL)
            self.floorEntry.configure(state=tkinter.NORMAL)
            self.sequenceEntry.configure(state=tkinter.NORMAL)

    def on_closing(self, event=0):
        self.destroy()

    def singleListReducer(self, lst, featuresLength):
        """
        Takes a single list and reduces it dimensions. This function is meant to be used
        in conjuction with the function below called "tupleDimensionReducer".
        """
        reducedList = []
        for i in range(featuresLength - 1):
            appender = float(lst[i] * 2 / 3) + float(lst[i + 1] * 1 / 3)
            reducedList.append(appender)
        return reducedList

    # print(singleListReducer(lstt, featureLength))

    def tupleDimensionReducer(self, tupleInput, featuresLength):
        """
        This function takes a tuple and reduces its dimensionality by 1.
        It does this by taking a "weighted average" between 2 dimensions and then adding them together.
        The first element in the tuple weights 2/3 while the second element weights 1/3.
        That pattern will continue trickling down until the last dimension has been reached.
        """

        # First we change the tuples into variables of type list. This makes us able to edit elements in them.
        for i in range(len(tupleInput)):
            tupleInput[i] = list(tupleInput[i])

        for i in range(len(tupleInput)):
            tupleInput[i] = self.singleListReducer(tupleInput[i], featuresLength)
        for i in range(len(tupleInput)):
            tupleInput[i] = tuple(tupleInput[i])
        return tupleInput

    def sequencerBM(self, propertyList):
        """
        This function takes a list of tuples, sorts them
        and then arranges them to alternate between lowest and highest.
        This ensures that BM avoids putting two complex modules right after each other.
        """
        listSorted = sorted(propertyList)

        endList = []
        for i in range(len(listSorted)):
            if "End" in listSorted[i]:
                endList.append(listSorted[i])
                listSorted.pop(i)
                break

        listRevSorted = []

        # The following code sorts the list of tuples alternating with the smallest value and then the higest value,
        # starting from the lowest value.
        for i in range(len(listSorted)):
            if (i % 2) == 0:
                # listRevSorted.append(2)
                listRevSorted.append(listSorted[-1])
                listSorted.pop(-1)
            elif (i % 2) == 1:
                # listRevSorted.append(1)
                listRevSorted.append(listSorted[0])
                listSorted.pop(0)
        for i in range(len(endList)):
            listRevSorted.append(endList[i])

        return listRevSorted

    def button_event2(self):
        results = [(1, 1, 5, 4, 2, "m1", "End"), (4, 1, 5, 4, 2, "m2"), (4, 1, 5, 4, 2, "m2"), (4, 1, 5, 4, 2, "m2"),
                   (2, 4, 4, 4, 2, "m3"), (2, 4, 4, 4, 2, "m3"), (2, 4, 4, 4, 2, "m3"), (2, 4, 4, 4, 2, "m3"),
                   (2, 4, 4, 4, 2, "m3"), (2, 4, 4, 4, 2, "m3")]
        results = self.sequencerBM(total)
        document = Document()
        for i in range(len(results)):
            self.sequenceField.configure(state='normal')
            self.sequenceField.insert(tkinter.END,
                                      str(results[i][:len(resolution)]) + '  ->  ' + str(results[i][6]) + '\n')
            self.sequenceField.configure(state='disabled')
            # print(results[i])
            if schedule == 1:
                self.documentation(results[i], i, document)

        #if schedule == 1:
        #    document.save("bm_plan.docx")
        #    convert("bm_plan.docx")

    def getProductionLength(self, leftScalarVector):
        leftScaled = [0, 0, 0, 0, 0, 0]

        # kitchen = 1
        if leftScalarVector[1] == 0:
            leftScaled[1] = 1
        else:
            leftScaled[1] = 2

        # Stairs = 2
        leftScaled[2] = leftScalarVector[2]

        # Ceiling, Windows & Doors = 3
        leftScaled[3] = leftScalarVector[3]

        # Outer shell = 4
        leftScaled[4] = 1

        # QC and Wrapping = 5
        leftScaled[5] = 1

        # Flooring & Scurting Boards = 0
        leftScaled[0] = 9 - sum(leftScaled)

        totalVector = np.concatenate(([2, 4, 3], leftScaled), axis=None)

        return totalVector

    def documentation(self, results, days, document):
        global weekends
        document.add_heading('BM Byggeindustri A/S - Ordrestyring', 0)

        p = document.add_paragraph(
            'Denne plan er vejledende ift. produktionstider og stationsstyring. Sedlen påsættes det nedenfor angivede modul.')

        # sel_date = sel_date + datetime.timedelta(days=weekends)
        # print(weekends)
        # print(sel_date)

        # production_length = [2, 3, 3, 2, 0, 2, 2, 2, 1]

        production_length = [1, 3, 3, 2, 0, 2, 2, 2, 1]
        production_length = self.getProductionLength(results)

        tom_values = [2, 5, 8, 10, 10, 12, 14, 16, 17]
        production_stations = ['Gips (t.o.m. ' + str(tom_values[0]) + ')',
                               'Spartel (t.o.m. ' + str(tom_values[1]) + ')',
                               'Maling (t.o.m. ' + str(tom_values[2]) + ')',
                               'Gulvlægning (t.o.m. ' + str(tom_values[3]) + ')',
                               'Trappe (t.o.m. ' + str(tom_values[4]) + ')',
                               'Indre installationer (t.o.m. ' + str(tom_values[5]) + ')',
                               'Ydre installationer (t.o.m. ' + str(tom_values[6]) + ')',
                               'Forskalling (t.o.m. ' + str(tom_values[7]) + ')',
                               'Kvalitetskontrol (t.o.m. ' + str(tom_values[8]) + ')']

        # Define date sets

        sel_date = self.cal.get_date()
        sel_date = sel_date + datetime.timedelta(days=days + weekends)
        while sel_date.weekday() in {5, 6}:
            weekends = weekends + 1
            sel_date = sel_date + datetime.timedelta(days=1)

        # print(sel_date)

        starting_date = sel_date

        production_records = []
        starting_dates = []

        weekends2 = 0

        for length in range(len(production_length)):
            production_dates = []
            lock = 0
            # starting_dates.append(sel_date + datetime.timedelta(days=1))
            for day in range(production_length[length]):
                added_date = sel_date + datetime.timedelta(days=1)
                if added_date.weekday() == 5:
                    weekends2 = weekends2 + 1
                    added_date = added_date + datetime.timedelta(days=2)
                sel_date = added_date
                if lock == 0:
                    starting_dates.append(sel_date)
                    lock = 1
            ## Check if any times are zero

            # print(added_date)

            date = (format_datetime(datetime=added_date, format='EEEE, d, MMMM, y', locale='da_DK'))

            production_dates.append(length)
            production_dates.append(production_stations[length])
            production_dates.append(date)
            production_dates.append('')

            production_records.append(production_dates)

        starting_dates.pop(0)
        starting_dates.insert(0, starting_date)

        for prod in range(len(production_length)):
            if production_length[prod] == 0:
                starting_dates.insert(prod, starting_dates[prod - 1])

        # print(production_records)
        # print(starting_dates)

        '''
        for dates in range(len(production_length)):




            production_dates = []
            added_date = sel_date + datetime.timedelta(days=production_length[dates])
            if added_date.weekday() + production_length[dates] >= 5:
                added_date = added_date + datetime.timedelta(days=2)
                #extra = extra + 2
            #added_date = added_date + datetime.timedelta(days=extra)
            #print(added_date.weekday())
            print(added_date)

            date = (format_datetime(datetime=added_date, format='EEEE, d, MMMM, y', locale='da_DK'))
            sel_date = added_date

            production_dates.append(production_stations[dates])
            production_dates.append(date)
            production_dates.append('')

            production_records.append(production_dates)

       # print(production_records)
        '''

        string = ('Denne seddel tilhører modul: ' + str(results[6]))

        moduleText = document.add_paragraph()
        moduleText.add_run(string).bold = True

        string = ('Opstart på modul: ' + str(
            (format_datetime(datetime=starting_date, format='EEEE, d, MMMM, y', locale='da_DK'))).capitalize())

        moduleText = document.add_paragraph()
        moduleText.add_run(string).bold = True
        # run = moduleText.add_run()
        # run.add_break()

        style = document.styles['Normal']
        font = style.font
        font.name = "Liberation Sans"
        font.size = Pt(13)

        table = document.add_table(rows=1, cols=4, style='Table Grid')
        row = table.rows[0]
        grade_formatted = row.cells[0].paragraphs[0].add_run('Frem til station:')
        grade_formatted.bold = True
        grade_formatted = row.cells[1].paragraphs[0].add_run('Opstartes d.')
        grade_formatted.bold = True
        grade_formatted = row.cells[2].paragraphs[0].add_run('Afsluttes d.')
        grade_formatted.bold = True
        grade_formatted = row.cells[3].paragraphs[0].add_run('Initialer')
        grade_formatted.bold = True

        for i, qty, id, desc in production_records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = str(
                (format_datetime(datetime=starting_dates[i], format='EEEE, d, MMMM, y', locale='da_DK'))).capitalize()
            row_cells[2].text = str(id).capitalize()
            row_cells[3].text = desc

        for row in table.rows:
            row.height = Cm(1.5)
        string = str(results[6][0]) + '_' + str(results[6][1]) + '_' + str(results[6][2])
        string = ('bmplan_' + string + '.docx')
        # print(string)
        document.add_page_break()


if __name__ == "__main__":
    app = App()
    app.mainloop()
